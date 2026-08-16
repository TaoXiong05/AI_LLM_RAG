"""把上传的 PDF / Docx / Txt / 图片文件解析成文本。

- PDF：提取文本页；"提取不到文字"的页（扫描页）用 PyMuPDF 渲染成图片交给 OCR。
- DOCX：读取段落文字 + 内嵌图片的 OCR 文字。
- 图片（png/jpg/webp/bmp）：直接 OCR。
"""
import io
from dataclasses import dataclass
from pathlib import Path

import pymupdf
from docx import Document as DocxDocument
from pypdf import PdfReader

from backend.ocr import ocr_image

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg", ".webp", ".bmp"}

IMAGE_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
}


class UnsupportedFileTypeError(ValueError):
    pass


@dataclass
class ParsedDocument:
    source: str
    text: str


def _ocr_pdf_page(pdf_doc, page_index: int) -> str:
    """把 PDF 的第 page_index 页渲染成 PNG 再 OCR。"""
    page = pdf_doc.load_page(page_index)
    pix = page.get_pixmap(dpi=200)
    return ocr_image(pix.tobytes("png"), "image/png")


def parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts: list[str] = []
    pdf_doc = None
    try:
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                text_parts.append(text)
                continue
            # 该页提取不到文本（扫描页/纯图片页）-> 渲染后 OCR
            if pdf_doc is None:
                pdf_doc = pymupdf.open(stream=file_bytes, filetype="pdf")
            text_parts.append(_ocr_pdf_page(pdf_doc, i))
    finally:
        if pdf_doc is not None:
            pdf_doc.close()
    return ParsedDocument(source=filename, text="\n".join(text_parts))


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]

    # 内嵌图片：逐个 OCR 并把识别文字拼到段落后面
    for shape in doc.inline_shapes:
        try:
            img_text = ocr_image(shape.blob)
        except Exception:  # noqa: BLE001 - 单个图片失败不影响整个文档
            img_text = ""
        if img_text:
            parts.append(f"\n【图片内容】\n{img_text}")

    return ParsedDocument(source=filename, text="\n".join(parts))


def parse_txt(file_bytes: bytes, filename: str) -> ParsedDocument:
    text = file_bytes.decode("utf-8", errors="replace")
    return ParsedDocument(source=filename, text=text)


def parse_image(file_bytes: bytes, filename: str, mime: str) -> ParsedDocument:
    return ParsedDocument(source=filename, text=ocr_image(file_bytes, mime))


def parse_file(uploaded_file) -> ParsedDocument:
    """uploaded_file 只需鸭子类型满足 .name / .getvalue()（Streamlit UploadedFile 即可）。"""
    filename = uploaded_file.name
    suffix = Path(filename).suffix.lower()
    file_bytes = uploaded_file.getvalue()

    if suffix == ".pdf":
        return parse_pdf(file_bytes, filename)
    if suffix == ".docx":
        return parse_docx(file_bytes, filename)
    if suffix == ".txt":
        return parse_txt(file_bytes, filename)
    if suffix in IMAGE_MIME:
        return parse_image(file_bytes, filename, IMAGE_MIME[suffix])

    raise UnsupportedFileTypeError(f"不支持的文件类型：{suffix}")
